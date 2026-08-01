from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PAPER_DIR = Path(__file__).resolve().parent
ASSET_DIR = PAPER_DIR / "assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
SOURCE = PAPER_DIR / "manuscript.md"
OUTPUT = PAPER_DIR / "DeepBlock_Baihua_CVRP_量子优化竞赛论文.docx"

FIGURES = {
    "fig_quantum_candidate_quality.png": ROOT / "results" / "baihua_quantum_candidate_quality_20260801" / "fig_quantum_candidate_quality.png",
    "fig1_formal_seed_improvements.png": ROOT / "results" / "baihua_competition_package_20260801" / "figures" / "fig1_formal_seed_improvements.png",
    "fig3_candidate_budget_hit_rate.png": ROOT / "results" / "baihua_competition_package_20260801" / "figures" / "fig3_candidate_budget_hit_rate.png",
    "fig4_ablation.png": ROOT / "results" / "baihua_competition_package_20260801" / "figures" / "fig4_ablation.png",
    "fig5_hardware_intervention.png": ROOT / "results" / "baihua_competition_package_20260801" / "figures" / "fig5_hardware_intervention.png",
}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F6F9"
GREY = "5B6573"
BLACK = "20252B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def table_widths(cols: int) -> list[int]:
    presets = {
        2: [3600, 5760],
        3: [2600, 2200, 4560],
        4: [2500, 2000, 2000, 2860],
        5: [2200, 1550, 2100, 2200, 1310],
        6: [1120, 1540, 1540, 1540, 1540, 2080],
    }
    return presets.get(cols, [9360 // cols] * (cols - 1) + [9360 - (9360 // cols) * (cols - 1)])


def set_run_font(run, ascii_font="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_style_font(style, ascii_font="Calibri", east_asia="Microsoft YaHei", size=11, bold=False, color=BLACK) -> None:
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_bottom_border(paragraph, color="C9D4E2", size="8", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("—  ")
    set_run_font(run, size=9, color=GREY)
    add_field(run, "PAGE")
    tail = paragraph.add_run("  —")
    set_run_font(tail, size=9, color=GREY)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, bold=True, color=BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True
    h1.paragraph_format.widow_control = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, bold=True, color=BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    set_style_font(cap, size=9, color=GREY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True

    if "Equation" not in styles:
        eq = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        eq = styles["Equation"]
    set_style_font(eq, ascii_font="Cambria Math", east_asia="Microsoft YaHei", size=10.5, color=DARK_BLUE)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(6)
    eq.paragraph_format.space_after = Pt(8)
    eq.paragraph_format.keep_together = True

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    set_style_font(ref, size=9.2, color=BLACK)
    ref.paragraph_format.left_indent = Inches(0.22)
    ref.paragraph_format.first_line_indent = Inches(-0.22)
    ref.paragraph_format.space_after = Pt(4)
    ref.paragraph_format.line_spacing = 1.08
    ref.paragraph_format.keep_together = True
    ref.paragraph_format.widow_control = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("DEEPBLOCK  ·  BAIHUA QUANTUM-CVRP")
    set_run_font(r, size=8.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run("                                      技术论文 · 2026")
    set_run_font(r2, size=8.5, color=GREY)
    set_bottom_border(p)
    first_header = section.first_page_header
    if not first_header.paragraphs:
        first_header.add_paragraph()
    first_header.paragraphs[0].clear()
    add_page_number(section.footer.paragraphs[0])
    section.first_page_footer.paragraphs[0].clear()


def add_inline(paragraph, text: str, default_size=None, default_color=None) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True, color=default_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, ascii_font="Consolas", east_asia="Microsoft YaHei", size=9.2 if default_size is None else default_size, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=default_size, color=default_color)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = table_widths(cols)
    for i, row_vals in enumerate(rows):
        row = table.rows[i]
        for j in range(cols):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, row_vals[j] if j < len(row_vals) else "", default_size=8.8, default_color=BLACK)
            if i == 0:
                set_cell_shading(cell, PALE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            elif i % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet_or_number(doc: Document, text: str, numbered=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    marker = f"{add_bullet_or_number.number_index}." if numbered else "•"
    if numbered:
        add_bullet_or_number.number_index += 1
    run = p.add_run(marker + "\t")
    set_run_font(run, size=10.5, bold=numbered, color=BLUE)
    add_inline(p, text)


add_bullet_or_number.number_index = 1


def add_figure(doc: Document, filename: str, caption: str) -> None:
    path = ASSET_DIR / filename if filename == "fig0_method_overview.png" else FIGURES[filename]
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    for doc_pr in run._r.xpath(".//wp:docPr"):
        doc_pr.set("title", caption.split("。", 1)[0])
        doc_pr.set("descr", caption)
    cap = doc.add_paragraph(style="Figure Caption")
    add_inline(cap, caption, default_size=9, default_color=GREY)


def add_equation(doc: Document, equation: str) -> None:
    p = doc.add_paragraph(style="Equation")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F9FC")
    p_pr.append(shd)
    run = p.add_run(equation)
    set_run_font(run, ascii_font="Cambria Math", east_asia="Microsoft YaHei", size=10.5, color=DARK_BLUE)


def draw_method_overview() -> Path:
    out = ASSET_DIR / "fig0_method_overview.png"
    image = Image.new("RGB", (2816, 1056), "white")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 30)
    semibold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 29)
    title_font = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 42)
    note_font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 29)
    scale_x, scale_y = 220, 220

    def sx(x):
        return int(x * scale_x)

    def sy(y):
        return int((4.8 - y) * scale_y)

    def center_text(box, text, font, fill):
        left, top, right, bottom = box
        bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=7, align="center")
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text(((left + right - tw) / 2, (top + bottom - th) / 2 - 2), text, font=font, fill=fill, spacing=7, align="center")

    def arrow(a, b, color="#607D8B", dashed=False):
        x1, y1 = sx(a[0]), sy(a[1])
        x2, y2 = sx(b[0]), sy(b[1])
        if dashed:
            segs = 9
            for idx in range(0, segs, 2):
                t1, t2 = idx / segs, min((idx + 1) / segs, 1)
                draw.line((x1 + (x2 - x1) * t1, y1 + (y2 - y1) * t1, x1 + (x2 - x1) * t2, y1 + (y2 - y1) * t2), fill=color, width=4)
        else:
            draw.line((x1, y1, x2, y2), fill=color, width=4)
        angle = math.atan2(y2 - y1, x2 - x1)
        length = 18
        wing = 0.55
        p1 = (x2 - length * math.cos(angle - wing), y2 - length * math.sin(angle - wing))
        p2 = (x2 - length * math.cos(angle + wing), y2 - length * math.sin(angle + wing))
        draw.polygon([(x2, y2), p1, p2], fill=color)

    nodes = [
        (0.25, 3.05, 2.15, 0.85, "CVRP input\n40 customers / 4 vehicles", "#EAF2F8", "#2E74B5"),
        (2.65, 3.05, 2.15, 0.85, "Classical feasible\ninitial solution", "#F4F6F9", "#6B7785"),
        (5.05, 3.05, 2.15, 0.85, "16-customer boundary pool\n3 overlapping blocks", "#EAF2F8", "#2E74B5"),
        (7.45, 3.05, 2.15, 0.85, "Topology-aligned\n8-bit sparse QUBO", "#E8F6F3", "#1E806C"),
        (9.85, 3.05, 2.55, 0.85, "Baihua QAOA sampling\np=1 · 4,096 shots", "#FFF3E6", "#C96A18"),
        (9.85, 1.15, 2.55, 0.85, "Independent quality gate\nbottom-10% energy mass", "#F4ECF7", "#7D3C98"),
        (7.45, 1.15, 2.15, 0.85, "Top-64 candidates\ncapacity repair", "#F4F6F9", "#6B7785"),
        (5.05, 1.15, 2.15, 0.85, "True route distance\nfull objective", "#F4F6F9", "#6B7785"),
        (2.65, 1.15, 2.15, 0.85, "Strict improvement?\nmonotone acceptance", "#E8F6F3", "#1E806C"),
        (0.25, 1.15, 2.15, 0.85, "Update solution\nand scan next block", "#EAF2F8", "#2E74B5"),
    ]
    for x, y, w, h, label, face, edge in nodes:
        box = (sx(x), sy(y + h), sx(x + w), sy(y))
        draw.rounded_rectangle(box, radius=18, fill=face, outline=edge, width=4)
        center_text(box, label, semibold, "#24313D")
    arrows = [
        ((2.40, 3.48), (2.65, 3.48)), ((4.80, 3.48), (5.05, 3.48)), ((7.20, 3.48), (7.45, 3.48)),
        ((9.60, 3.48), (9.85, 3.48)), ((11.13, 3.05), (11.13, 2.00)), ((9.85, 1.58), (9.60, 1.58)),
        ((7.45, 1.58), (7.20, 1.58)), ((5.05, 1.58), (4.80, 1.58)), ((2.65, 1.58), (2.40, 1.58)),
    ]
    for a, b in arrows:
        arrow(a, b)
    # Dashed loop-back arrow beside the first column.
    draw.arc((sx(0.52), sy(3.05), sx(2.12), sy(1.15)), start=90, end=270, fill="#2E74B5", width=4)
    arrow((0.52, 2.15), (0.52, 3.02), color="#2E74B5", dashed=True)
    title = "DeepBlock topology-co-designed hybrid refinement loop"
    tb = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((2816 - (tb[2] - tb[0])) / 2, 64), title, font=title_font, fill="#1F4D78")
    note = "Quantum: candidate generation  |  Classical: feasibility, objective evaluation and acceptance"
    nb = draw.textbbox((0, 0), note, font=note_font)
    draw.text(((2816 - (nb[2] - nb[0])) / 2, 954), note, font=note_font, fill="#5B6573")
    image.save(out, dpi=(220, 220))
    return out


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("量子 + 优化竞赛技术论文")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("DeepBlock：面向容量约束车辆路径问题的\n拓扑协同QAOA真机混合精化")
    set_run_font(r, size=25, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Topology-Co-Designed QAOA Hardware Refinement for the Capacitated Vehicle Routing Problem")
    set_run_font(r, size=12, color=GREY, italic=True)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    vals = [
        ("2.73×", "真机低能候选富集"),
        ("147,456", "Baihua总shots"),
        ("36 / 36", "正式任务0 SWAP"),
    ]
    for i, (value, label) in enumerate(vals):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, LIGHT_BLUE if i != 1 else "F4F6F9")
        set_cell_margins(cell, top=180, start=100, bottom=180, end=100)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(value)
        set_run_font(r1, size=18, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, color=GREY)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [3120, 3120, 3120])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("作者：参赛团队（姓名待补）")
    set_run_font(r, size=11, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("单位：待补  ·  联系方式：待补")
    set_run_font(r, size=10, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("结论边界")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.6)
    p.paragraph_format.right_indent = Inches(0.6)
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("真机显著富集低BQM能量候选，并在混合闭环中产生有效精化；现有证据不支持相对经典随机或工业求解器的量子加速。")
    set_run_font(r, size=10, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("2026年8月  ·  完整实验论文版")
    set_run_font(r, size=9.5, color=GREY)
    p.add_run().add_break(WD_BREAK.PAGE)


def parse_manuscript(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    in_references = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line == "[[PAGEBREAK]]":
            doc.add_page_break()
            add_bullet_or_number.number_index = 1
            i += 1
            continue
        fig_match = re.fullmatch(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", line)
        if fig_match:
            add_figure(doc, fig_match.group(1), fig_match.group(2))
            i += 1
            continue
        eq_match = re.fullmatch(r"\[\[EQUATION:(.+)\]\]", line)
        if eq_match:
            add_equation(doc, eq_match.group(1))
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[3:])
            add_bullet_or_number.number_index = 1
            i += 1
            continue
        if line.startswith("# "):
            text = line[2:]
            p = doc.add_paragraph(style="Heading 1")
            if text == "参考文献":
                p.paragraph_format.page_break_before = True
            add_inline(p, text)
            in_references = text == "参考文献"
            add_bullet_or_number.number_index = 1
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        num = re.match(r"^(\d+)\.\s+(.+)$", line)
        if num and not in_references:
            add_bullet_or_number.number_index = int(num.group(1))
            add_bullet_or_number(doc, num.group(2), numbered=True)
            i += 1
            continue
        if line.startswith("- "):
            add_bullet_or_number(doc, line[2:], numbered=False)
            i += 1
            continue
        if in_references and re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph(style="Reference")
            add_inline(p, line, default_size=9.2)
            i += 1
            continue

        # Join wrapped plain-text lines into one semantic paragraph.
        chunks = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "- ", "[[")) or re.match(r"^\d+\.\s+", nxt):
                break
            chunks.append(nxt)
            i += 1
        paragraph_text = " ".join(chunks)
        p = doc.add_paragraph()
        add_inline(p, paragraph_text)


def build() -> Path:
    draw_method_overview()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "DeepBlock：面向容量约束车辆路径问题的拓扑协同QAOA真机混合精化"
    doc.core_properties.subject = "量子+优化竞赛完整实验论文"
    doc.core_properties.author = "参赛团队（姓名待补）"
    doc.core_properties.keywords = "CVRP, QAOA, QUBO, Baihua, hybrid quantum-classical optimization"
    add_title_page(doc)
    parse_manuscript(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
